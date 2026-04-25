# backend/pipeline/report_generator.py
"""
Генератор DOCX-отчёта.
Принципы:
  - нейтральные формулировки (факты, отклонения, возможные причины)
  - никаких обвинительных выводов без достаточной базы
  - каждый сигнал подкреплён статистикой (медиана, N наблюдений)
  - fallback без AI если API недоступен
"""

import io
import os
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# ─── Цвета ────────────────────────────────────────────────────────

COLOR_CRITICAL = RGBColor(0x7C, 0x3A, 0xED)
COLOR_HIGH     = RGBColor(0xCC, 0x33, 0x33)
COLOR_MEDIUM   = RGBColor(0xD4, 0x88, 0x2A)
COLOR_LOW      = RGBColor(0x2E, 0xA8, 0x4A)
COLOR_ACCENT   = RGBColor(0x1F, 0x49, 0x7D)

RISK_COLORS = {"CRITICAL": COLOR_CRITICAL, "HIGH": COLOR_HIGH, "MEDIUM": COLOR_MEDIUM, "LOW": COLOR_LOW}
RISK_LABELS = {"CRITICAL": "КРИТИЧЕСКИЙ", "HIGH": "ВЫСОКИЙ", "MEDIUM": "СРЕДНИЙ",    "LOW": "НИЗКИЙ"}
RISK_BG     = {"CRITICAL": "EDE7F6",      "HIGH": "FFEBEE",   "MEDIUM": "FFF8E1",    "LOW": "F1F8E9"}

SCORE_EXPLANATION = (
    "Скор рассчитывается аддитивно: каждый активный индикатор добавляет "
    "фиксированное количество баллов (сумма ограничена 100). "
    "Скор отражает количество и значимость сигналов — не вероятность нарушения."
)


# ─── Нарратив (детерминированный, без AI) ─────────────────────────

def _build_narrative(results: list, source_files: list) -> str:
    critical = [r for r in results if r.get("risk_level") == "CRITICAL"]
    high     = [r for r in results if r.get("risk_level") == "HIGH"]
    medium   = [r for r in results if r.get("risk_level") == "MEDIUM"]
    low      = [r for r in results if r.get("risk_level") == "LOW"]
    total    = len(results)
    total_sum = sum(r.get("total_price") or 0 for r in results)

    lines = []

    # ── Резюме ──
    lines.append("## ИСПОЛНИТЕЛЬНОЕ РЕЗЮМЕ")
    attention = len(critical) + len(high)
    lines.append(
        f"Проанализировано {total} позиций из {len(source_files or [])} документов. "
        f"Общий объём закупок: {total_sum:,.0f} руб. "
        f"Выявлено позиций, требующих внимания: {attention} "
        f"(критический уровень: {len(critical)}, высокий: {len(high)}, средний: {len(medium)})."
    )

    if attention == 0:
        lines.append(
            "По результатам автоматического анализа позиций с высоким или критическим "
            "уровнем индикаторов не выявлено. Рекомендуется плановый мониторинг."
        )
    else:
        lines.append(
            f"Позиции с ненулевым скором содержат один или несколько индикаторов отклонения: "
            f"расхождение цен между документами, несоответствие объёмов, неполнота данных "
            f"или структурные особенности закупки. "
            f"Каждый сигнал описан в разделе «Детальный разбор». "
            f"Выводы носят индикативный характер и требуют верификации по первичным документам."
        )

    # ── Детали ──
    lines.append("\n## ДЕТАЛЬНЫЙ АНАЛИЗ ИНДИКАТОРОВ")

    for r in sorted(critical + high, key=lambda x: x.get("score", 0), reverse=True)[:10]:
        name  = r.get("item") or r.get("name") or "—"
        score = r.get("score", 0)
        rl    = r.get("risk_level", "")
        full  = r.get("full_explanation", {})

        lines.append(f"\n### {name}  (скор: {score}/100, уровень: {RISK_LABELS.get(rl, rl)})")

        flags_explained = full.get("flags_explained", [])
        if flags_explained:
            for fe in flags_explained:
                facts  = fe.get("facts", "")
                dev    = fe.get("deviation", "")
                interp = fe.get("interpretation", "")
                if facts:
                    lines.append(f"ФАКТЫ: {facts}")
                if dev and dev != "—":
                    lines.append(f"ОТКЛОНЕНИЕ: {dev}")
                if interp:
                    lines.append(f"ИНТЕРПРЕТАЦИЯ: {interp}")
                lines.append("")
        else:
            expl = r.get("explanation", "")
            if expl:
                lines.append(expl)

    # ── Методология скоринга ──
    lines.append("\n## МЕТОДОЛОГИЯ СКОРИНГА")
    lines.append(SCORE_EXPLANATION)
    lines.append("\nУровни риска:")
    lines.append("— CRITICAL (70–100): 3+ значимых индикатора одновременно")
    lines.append("— HIGH (40–69): несколько индикаторов или один весомый")
    lines.append("— MEDIUM (20–39): один-два слабых сигнала")
    lines.append("— LOW (0–19): единичный незначительный сигнал или его отсутствие")

    # ── Рекомендации ──
    lines.append("\n## РЕКОМЕНДУЕМЫЕ ДЕЙСТВИЯ")
    if critical:
        lines.append("**По позициям с уровнем CRITICAL:**")
        lines.append("— Запросить первичные документы: договоры, счета-фактуры, акты выполненных работ.")
        lines.append("— Сверить данные с бухгалтерским учётом и банковскими выписками.")
        lines.append("— Получить письменное обоснование от ответственного сотрудника.")
        lines.append("")
    if high:
        lines.append("**По позициям с уровнем HIGH:**")
        lines.append("— Включить в ближайшую плановую проверку.")
        lines.append("— Запросить актуальный прайс-лист поставщика для сверки цен.")
        lines.append("")
    lines.append("**Общее:**")
    lines.append("— Результаты автоматического анализа не являются основанием для выводов о нарушениях.")
    lines.append("— Окончательное заключение формируется по итогам документальной проверки.")

    # ── Выводы ──
    lines.append("\n## ВЫВОДЫ")
    lines.append(
        f"Автоматический анализ выявил {attention} позиций с индикаторами отклонения "
        f"из {total} проанализированных. "
        f"Система фиксирует статистические отклонения и структурные несоответствия — "
        f"интерпретация причин остаётся за ответственным специалистом."
    )

    return "\n".join(lines)


def _try_ai_narrative(results: list, source_files: list) -> str | None:
    """Пробует Anthropic API, возвращает None если недоступен."""
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return None
    try:
        import anthropic
        high_risk = [r for r in results if r.get("score", 0) >= 40]
        items_summary = []
        for r in high_risk[:10]:
            full  = r.get("full_explanation", {})
            flags = "; ".join(
                f"{fe['deviation']}" for fe in full.get("flags_explained", [])
                if fe.get("deviation") and fe["deviation"] != "—"
            )
            items_summary.append(
                f'- «{r.get("name","?")}»: скор {r.get("score")}/100, индикаторы: {flags or r.get("explanation","")}'
            )

        prompt = f"""Ты — аналитик внутреннего контроля. Составь нейтральное аналитическое резюме.

Документов: {len(source_files or [])}
Позиций: {len(results)}
Позиций с индикаторами (скор ≥40): {len(high_risk)}

Ключевые позиции:
{chr(10).join(items_summary) if items_summary else '— отклонений не выявлено'}

Требования к тексту:
1. Только нейтральные формулировки: «выявлено отклонение», «требует проверки», «возможные причины»
2. Никаких обвинительных выводов
3. Каждый вывод подкреплён конкретными числами из данных
4. Структура: РЕЗЮМЕ → ДЕТАЛИ → РЕКОМЕНДАЦИИ → ВЫВОДЫ
5. Деловой язык, без воды"""

        client = anthropic.Anthropic(api_key=api_key)
        msg    = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=3000,
            messages=[{"role": "user", "content": prompt}],
        )
        return msg.content[0].text.strip()
    except Exception as e:
        logger.warning(f"AI narrative failed: {e}")
        return None


def _generate_narrative(results: list, source_files: list) -> str:
    ai = _try_ai_narrative(results, source_files)
    if ai:
        return ai
    return _build_narrative(results, source_files)


# ─── DOCX утилиты ─────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, color="CCCCCC"):
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0");    b.set(qn("w:color"), color)
        tcBdr.append(b)
    tcPr.append(tcBdr)


def _set_col_width(cell, dxa: int):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(dxa)); tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _para(doc, text, bold=False, size=11, color=None,
          align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6, italic=False):
    p   = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold       = bold
    run.italic     = italic
    run.font.size  = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def _heading(doc, text, level=1):
    sizes = {1: 13, 2: 11, 3: 10}
    p = _para(doc, text, bold=True, size=sizes.get(level, 10),
              color=COLOR_ACCENT, space_before=10, space_after=4)
    if level == 1:
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
        bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), "1F497D")
        pBdr.append(bot); pPr.append(pBdr)
    return p


def _render_narrative(doc, narrative: str):
    for line in narrative.splitlines():
        line = line.rstrip()
        if not line:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        if line.startswith("## "):
            _heading(doc, line[3:], level=1)
        elif line.startswith("### "):
            _heading(doc, line[4:], level=2)
        elif line.startswith("**") and line.endswith("**"):
            _para(doc, line.strip("*"), bold=True, size=9,
                  color=COLOR_ACCENT, space_before=4, space_after=2)
        elif line.startswith("— ") or line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            p.add_run(line[2:]).font.size = Pt(9)
        elif line.startswith("ФАКТЫ:"):
            _para(doc, line, bold=False, size=9,
                  color=RGBColor(0x20, 0x40, 0x70), space_before=2, space_after=1)
        elif line.startswith("ОТКЛОНЕНИЕ:"):
            _para(doc, line, bold=True, size=9,
                  color=COLOR_HIGH, space_before=1, space_after=1)
        elif line.startswith("ИНТЕРПРЕТАЦИЯ:"):
            _para(doc, line, italic=True, size=9,
                  color=RGBColor(0x50, 0x50, 0x50), space_before=1, space_after=4)
        else:
            _para(doc, line, size=9, space_before=0, space_after=3)


# ─── Сводная таблица ──────────────────────────────────────────────

def _render_summary_table(doc, results: list):
    total    = len(results)
    critical = sum(1 for r in results if r.get("risk_level") == "CRITICAL")
    high     = sum(1 for r in results if r.get("risk_level") == "HIGH")
    medium   = sum(1 for r in results if r.get("risk_level") == "MEDIUM")
    low      = sum(1 for r in results if r.get("risk_level") == "LOW")
    avg      = round(sum(r.get("score", 0) for r in results) / total, 1) if total else 0

    tbl = doc.add_table(rows=2, cols=6)
    tbl.style = "Table Grid"
    hdrs = ["Позиций", "CRITICAL", "HIGH", "MEDIUM", "LOW", "Средний скор"]
    vals = [str(total), str(critical), str(high), str(medium), str(low), f"{avg}/100"]
    vcols = {"CRITICAL": COLOR_CRITICAL, "HIGH": COLOR_HIGH, "MEDIUM": COLOR_MEDIUM}

    for i, (h, v) in enumerate(zip(hdrs, vals)):
        hc = tbl.rows[0].cells[i]
        vc = tbl.rows[1].cells[i]
        _set_cell_bg(hc, "1F497D"); _set_cell_border(hc, "1F497D"); _set_cell_border(vc)
        hc.paragraphs[0].clear()
        hr = hc.paragraphs[0].add_run(h)
        hr.font.size = Pt(9); hr.font.bold = True
        hr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        vc.paragraphs[0].clear()
        vr = vc.paragraphs[0].add_run(v)
        vr.font.size = Pt(12); vr.font.bold = True
        if h in vcols and int(v) > 0:
            vr.font.color.rgb = vcols[h]
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


# ─── Таблица позиций ──────────────────────────────────────────────

def _render_items_table(doc, results: list):
    sorted_r = sorted(results, key=lambda r: r.get("score", 0), reverse=True)
    widths   = [4800, 1200, 1400, 7000]

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"

    hrow  = tbl.rows[0]
    hhdrs = ["Позиция / Поставщик", "Скор", "Уровень", "Факты и отклонения"]
    for i, (cell, hdr, w) in enumerate(zip(hrow.cells, hhdrs, widths)):
        _set_cell_bg(cell, "1F497D"); _set_cell_border(cell, "1F497D"); _set_col_width(cell, w)
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(hdr)
        r.font.size = Pt(9); r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, result in enumerate(sorted_r):
        rl    = result.get("risk_level", "LOW")
        score = result.get("score", 0)
        name  = result.get("item") or result.get("name") or "—"
        conts = result.get("contractors") or []
        tp    = result.get("total_price") or 0

        row = tbl.add_row()
        bg  = RISK_BG.get(rl, "FFFFFF") if score >= 20 else ("F8F8F8" if idx % 2 else "FFFFFF")
        for cell, w in zip(row.cells, widths):
            _set_cell_border(cell); _set_cell_bg(cell, bg); _set_col_width(cell, w)

        # Колонка 0: название
        c0 = row.cells[0]
        c0.paragraphs[0].clear()
        r0 = c0.paragraphs[0].add_run(name)
        r0.font.size = Pt(9); r0.font.bold = True
        if tp > 0:
            p2 = c0.add_paragraph()
            p2.paragraph_format.space_before = Pt(1)
            p2.add_run(f"Сумма: {tp:,.0f} руб.").font.size = Pt(8)
        if conts:
            p3 = c0.add_paragraph()
            p3.add_run(conts[0][:55] + ("…" if len(conts[0]) > 55 else "")).font.size = Pt(7.5)

        # Колонка 1: скор
        c1 = row.cells[1]
        c1.paragraphs[0].clear()
        rs = c1.paragraphs[0].add_run(str(score))
        rs.font.size = Pt(14); rs.font.bold = True
        rs.font.color.rgb = RISK_COLORS.get(rl, COLOR_LOW)
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Колонка 2: уровень
        c2 = row.cells[2]
        c2.paragraphs[0].clear()
        rr = c2.paragraphs[0].add_run(RISK_LABELS.get(rl, rl))
        rr.font.size = Pt(8); rr.font.bold = True
        rr.font.color.rgb = RISK_COLORS.get(rl, COLOR_LOW)
        c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Колонка 3: факты и отклонения из full_explanation
        c3   = row.cells[3]
        c3.paragraphs[0].clear()
        full = result.get("full_explanation", {})
        flags_explained = full.get("flags_explained", [])

        if flags_explained:
            first = True
            for fe in flags_explained:
                facts  = fe.get("facts", "")
                dev    = fe.get("deviation", "")
                interp = fe.get("interpretation", "")

                # ФАКТЫ
                if facts:
                    p = c3.paragraphs[0] if first else c3.add_paragraph()
                    p.paragraph_format.space_before = Pt(0 if first else 4)
                    rf = p.add_run(f"ФАКТЫ: {facts}")
                    rf.font.size = Pt(8)
                    rf.font.color.rgb = RGBColor(0x20, 0x40, 0x70)
                    first = False

                # ОТКЛОНЕНИЕ
                if dev and dev != "—":
                    pd = c3.add_paragraph()
                    pd.paragraph_format.space_before = Pt(1)
                    rd = pd.add_run(f"ОТКЛОНЕНИЕ: {dev}")
                    rd.font.size = Pt(8); rd.font.bold = True
                    rd.font.color.rgb = RISK_COLORS.get(rl, COLOR_MEDIUM)

                # ИНТЕРПРЕТАЦИЯ
                if interp:
                    pi = c3.add_paragraph()
                    pi.paragraph_format.space_before = Pt(1)
                    pi.paragraph_format.space_after  = Pt(4)
                    ri = pi.add_run(f"Возможные причины: {interp.replace('Возможные причины: ', '')}")
                    ri.font.size = Pt(7.5); ri.italic = True
                    ri.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        else:
            expl = result.get("explanation", "Отклонений не выявлено")
            c3.paragraphs[0].add_run(expl).font.size = Pt(8.5)


# ─── Основная функция ─────────────────────────────────────────────

def generate_report(results: list, source_files: list = None) -> bytes:
    doc = Document()

    section = doc.sections[0]
    section.page_width    = Cm(21);  section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.5); section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0); section.bottom_margin = Cm(2.0)

    date_str = datetime.now().strftime("%d.%m.%Y %H:%M")

    # ── Шапка ──
    _para(doc, "АНАЛИТИЧЕСКИЙ ОТЧЁТ", bold=True, size=18,
          color=COLOR_ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER,
          space_before=0, space_after=2)
    _para(doc, "Автоматический анализ финансовых документов на предмет отклонений",
          size=10, color=RGBColor(0x60, 0x60, 0x60),
          align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2)
    _para(doc, f"Сформирован: {date_str}", size=9,
          color=RGBColor(0x80, 0x80, 0x80),
          align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=10)

    # Дисклеймер
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(
        "⚠  Данный отчёт сформирован автоматически на основе статистического анализа данных. "
        "Все выявленные позиции являются индикаторами и требуют верификации по первичным документам. "
        "Система не устанавливает факт нарушений и не делает обвинительных выводов."
    )
    r.font.size = Pt(8.5); r.italic = True
    r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    # ── 1. Документы ──
    _heading(doc, "1. Проанализированные документы", level=1)
    if source_files:
        for f in source_files:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            p.add_run(f).font.size = Pt(10)
    else:
        _para(doc, "Список файлов не передан.", size=10)

    # ── 2. Сводка ──
    _heading(doc, "2. Сводная статистика", level=1)
    _render_summary_table(doc, results)
    doc.add_paragraph()

    # Методология скоринга
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"Методология: {SCORE_EXPLANATION}")
    r.font.size = Pt(8); r.italic = True
    r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    # ── 3. Аналитическое заключение ──
    _heading(doc, "3. Аналитическое заключение", level=1)
    narrative = _generate_narrative(results, source_files or [])
    _render_narrative(doc, narrative)
    doc.add_paragraph()

    # ── 4. Детальный разбор ──
    _heading(doc, "4. Детальный разбор позиций", level=1)
    _para(doc,
          "Для каждой позиции указаны: ФАКТЫ (конкретные числа), "
          "ОТКЛОНЕНИЕ (отклонение от медианы с базой), "
          "ВОЗМОЖНЫЕ ПРИЧИНЫ (нейтральная интерпретация).",
          size=8.5, italic=True, color=RGBColor(0x70, 0x70, 0x70), space_after=8)
    _render_items_table(doc, results)
    doc.add_paragraph()

    # ── Подпись ──
    _para(doc, "─" * 55, size=8, color=RGBColor(0xCC, 0xCC, 0xCC),
          space_before=14, space_after=3)
    _para(doc,
          f"Отчёт сформирован автоматически. {date_str}. "
          f"Данные обработаны локально. Требует верификации специалистом.",
          size=7.5, color=RGBColor(0x90, 0x90, 0x90))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()