# backend/pipeline/parser.py
"""
Универсальный парсер документов.

Стратегия для XLSX:
1. Извлекаем метаданные (заказчик, подрядчик, стройка) из шапки
2. Находим строку с номерами колонок (1,2,3...) — это маркер начала таблицы
3. Из строк-заголовков над ней определяем смысл каждой колонки по ключевым словам
4. Читаем строки данных, применяя найденный маппинг
5. Если таблица не найдена — fallback на плоский текст

Стратегия для PDF/DOCX: извлекаем структурированный текст + таблицы.
"""

import re
import pdfplumber
import docx
import openpyxl
from pathlib import Path
from typing import Any


# ─── Ключевые слова для определения колонок ───────────────────────

_COL_KEYWORDS = {
    "name": [
        "наименование", "название", "вид работ", "виды работ",
        "конструктивные элементы", "работ", "услуг", "материал",
        "товар", "позиция", "описание", "номенклатура",
    ],
    "quantity": [
        "количество", "кол-во", "кол.", "объём", "объем", "qty",
        "количе", "кол-ство",
    ],
    "unit": [
        "единица измерения", "ед. изм", "ед.изм", "еди- ница",
        "единица", "ед.", "unit",
    ],
    "unit_price": [
        "цена за единицу", "цена за ед", "расценка", "unit price",
        "цена за", "стоимость ед", "цена ед",
    ],
    "total_price": [
        "стоимость", "сумма", "итого", "итоговая", "total",
        "всего", "общая стоимость", "итог", "стои- мость",
        "стоимость выполненных", "стоимость фактически",
    ],
    "contractor": [
        "поставщик", "подрядчик", "исполнитель", "контрагент",
        "организация",
    ],
    "date": [
        "дата", "период", "месяц", "дата поставки",
    ],
}

# Ключевые слова для извлечения метаданных из шапки
_META_KEYWORDS = {
    "contractor": ["подрядчик", "субподрядчик", "исполнитель", "поставщик"],
    "customer":   ["заказчик", "генподрядчик", "инвестор"],
    "project":    ["стройка", "объект", "проект"],
}

_JUNK_ROWS = {
    "итого", "всего", "накладные", "прочие расходы", "итог",
    "в том числе", "х", "x", "сметная стоимость",
}

_UNIT_RE = re.compile(
    r'\b(шт\.?|кг|тонн[а-я]*|т\.?\b|м2|м3|кв\.?\s*м|куб\.?\s*м|'
    r'л\.?\b|метр[а-я]*|компл\.?|уп\.?|пог\.?\s*м|ч\.?\b|смена|'
    r'контейнер|рейс|км)\b',
    re.IGNORECASE
)


# ─── Утилиты ──────────────────────────────────────────────────────

def _to_float(x: Any) -> float | None:
    if x is None:
        return None
    s = str(x).strip().replace("\xa0", "").replace(" ", "")
    # Русский формат цены: '350-00', '7 000 000-00' → дефис как разделитель копеек
    s = re.sub(r"-\d{2}$", "", s)
    s = s.replace(",", ".")
    s = re.sub(r"[^\d.]", "", s)
    if not s:
        return None
    try:
        v = float(s)
        return v if v >= 0 else None
    except ValueError:
        return None


def _clean_str(x: Any) -> str:
    if x is None:
        return ""
    return re.sub(r"\s+", " ", str(x)).strip()


def _cell_text(x: Any) -> str:
    return _clean_str(x).lower()


def _is_junk_row(name: str) -> bool:
    n = name.lower().strip()
    return any(j in n for j in _JUNK_ROWS) or len(n) < 2


def _extract_unit(s: str) -> str:
    m = _unit_re_search(s)
    return m.group(0).strip().lower() if m else ""


def _unit_re_search(s: str):
    return _UNIT_RE.search(str(s))


# ─── XLSX парсер ──────────────────────────────────────────────────

def _score_col_keyword(cell_text: str, col_type: str) -> int:
    """Возвращает счёт совпадения ячейки с типом колонки."""
    score = 0
    for kw in _COL_KEYWORDS[col_type]:
        if kw in cell_text:
            score += len(kw)  # длиннее совпадение → точнее
    return score


def _detect_column_mapping(header_rows: list[list]) -> dict[str, int]:
    """
    Определяет маппинг {тип: индекс_колонки} из нескольких строк заголовка.
    Объединяет тексты всех заголовочных строк по колонкам.
    """
    if not header_rows:
        return {}

    n_cols = max(len(r) for r in header_rows)
    col_texts = [""] * n_cols
    for row in header_rows:
        for ci, cell in enumerate(row):
            t = _cell_text(cell)
            if t:
                col_texts[ci] = (col_texts[ci] + " " + t).strip()

    # Колонки с сервисным содержимым (номера, порядок) — исключаем из числовых типов
    _INDEX_KW = {"номер", "по поряд", "порядк", "поряд", "пози", "№"}
    index_cols = {
        ci for ci, t in enumerate(col_texts)
        if any(kw in t for kw in _INDEX_KW)
    }

    # Приоритет назначения: специфичные типы раньше общих
    PRIORITY = ["unit_price", "total_price", "quantity", "unit", "name", "contractor", "date"]

    # Строим матрицу scores[col_type][col_idx]
    score_matrix: dict[str, dict[int, int]] = {}
    for col_type in PRIORITY:
        score_matrix[col_type] = {}
        for ci, text in enumerate(col_texts):
            # Числовые колонки не назначаем индексным столбцам
            if col_type in ("unit_price", "total_price", "quantity") and ci in index_cols:
                continue
            s = _score_col_keyword(text, col_type)
            if s > 0:
                score_matrix[col_type][ci] = s

    mapping: dict[str, int] = {}
    used: set[int] = set()

    for col_type in PRIORITY:
        candidates = score_matrix[col_type]
        if not candidates:
            continue
        # Среди незанятых — выбираем с максимальным счётом,
        # при равных — правее (для числовых полей правые колонки вероятнее)
        free = {ci: s for ci, s in candidates.items() if ci not in used}
        if not free:
            # Разрешаем переназначение только если счёт значительно выше
            all_c = sorted(candidates.items(), key=lambda x: (-x[1], -x[0]))
            if all_c:
                mapping[col_type] = all_c[0][0]
            continue

        best = sorted(
            free.items(),
            key=lambda x: (-x[1], -x[0] if col_type in ("unit_price", "total_price") else x[0])
        )[0][0]
        mapping[col_type] = best
        used.add(best)

    return mapping


def _find_table_start(ws) -> tuple[int, list[list], dict]:
    """
    Ищет начало таблицы данных.
    Маркер: строка где первые ячейки содержат числа 1,2,3 (номера колонок).
    Заголовки — строки выше этой строки (но не дальше 10 строк выше).
    Возвращает: (data_start_row_idx, header_rows, col_mapping)
    """
    rows = list(ws.iter_rows(values_only=True))

    for i, row in enumerate(rows):
        vals = [v for v in row if v is not None]
        # Строка номеров колонок: содержит 1, 2, 3 как числа или строки
        nums = set()
        for v in vals[:8]:
            f = _to_float(v)
            if f is not None and 1 <= f <= 10:
                nums.add(int(f))
        if {1, 2, 3}.issubset(nums):
            # Нашли строку с номерами — данные начинаются со следующей
            data_start = i + 1
            # Заголовки — до 6 строк выше
            header_start = max(0, i - 6)
            header_rows  = [list(rows[j]) for j in range(header_start, i)]
            mapping      = _detect_column_mapping(header_rows)
            return data_start, header_rows, mapping

    # Fallback: ищем строку с ключевыми словами наименования
    for i, row in enumerate(rows[:30]):
        texts = [_cell_text(v) for v in row if v is not None]
        if any("наименование" in t or "вид работ" in t for t in texts):
            header_rows = [list(row)]
            mapping     = _detect_column_mapping(header_rows)
            return i + 1, header_rows, mapping

    return -1, [], {}


def _extract_metadata(ws) -> dict:
    """Извлекает метаданные (подрядчик, заказчик, объект) из шапки документа."""
    meta = {"contractor": "", "customer": "", "project": "", "date": ""}
    rows = list(ws.iter_rows(values_only=True))

    for row in rows[:35]:
        for ci, cell in enumerate(row):
            cell_low = _cell_text(cell)
            if not cell_low:
                continue
            for meta_key, keywords in _META_KEYWORDS.items():
                if any(kw in cell_low for kw in keywords):
                    # Ищем значение правее в этой строке
                    for offset in range(1, min(15, len(row) - ci)):
                        val = _clean_str(row[ci + offset] if ci + offset < len(row) else None)
                        if val and len(val) > 3 and "(организация" not in val.lower():
                            if not meta.get(meta_key):
                                meta[meta_key] = val
                            break

        # Дата из шапки
        for cell in row:
            s = _clean_str(cell)
            m = re.search(r'\b(\d{1,2}[./]\d{1,2}[./]\d{4})\b', s)
            if m and not meta["date"]:
                meta["date"] = m.group(1)

    return meta


def _parse_row_to_item(row: list, mapping: dict, meta: dict) -> dict | None:
    """Преобразует строку данных в item по маппингу."""

    def cell(key):
        i = mapping.get(key)
        return row[i] if i is not None and i < len(row) else None

    name = _clean_str(cell("name"))
    # Fallback: merged cells могут сдвигать name на соседний столбец
    if not name or _is_junk_row(name):
        for ci, val in enumerate(row):
            s = _clean_str(val)
            if len(s) >= 4 and _to_float(val) is None and ci not in (0, 1, 2):
                if not _is_junk_row(s):
                    name = s
                    break
    if not name or _is_junk_row(name):
        return None

    # Числовые поля
    quantity    = _to_float(cell("quantity"))
    unit_price  = _to_float(cell("unit_price"))
    total_price = _to_float(cell("total_price"))

    # Единица измерения — из колонки или из названия
    unit_raw = _clean_str(cell("unit"))
    unit     = _extract_unit(unit_raw) or _extract_unit(name)

    # Восстановление
    if unit_price and quantity and not total_price:
        total_price = round(unit_price * quantity, 2)
    if total_price and quantity and quantity > 0 and not unit_price:
        unit_price = round(total_price / quantity, 2)
    if total_price and unit_price and unit_price > 0 and not quantity:
        quantity = round(total_price / unit_price, 4)

    # Должна быть хоть какая-то цена
    if not unit_price and not total_price:
        return None

    contractor = _clean_str(cell("contractor")) or meta.get("contractor", "")
    date       = _clean_str(cell("date")) or meta.get("date", "")

    return {
        "name":        name,
        "quantity":    quantity,
        "unit":        unit,
        "unit_price":  unit_price,
        "total_price": total_price,
        "contractor":  contractor,
        "date":        date,
        "_has_detail": quantity is not None and unit_price is not None,
    }


def parse_xlsx(file_path: Path) -> dict:
    wb = openpyxl.load_workbook(file_path, data_only=True)
    all_items = []
    meta_global = {"contractor": "", "customer": "", "project": "", "date": ""}

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]

        meta = _extract_metadata(ws)
        # Берём первый найденный подрядчик как глобальный
        for key in meta_global:
            if not meta_global[key] and meta.get(key):
                meta_global[key] = meta[key]

        data_start, header_rows, mapping = _find_table_start(ws)
        if data_start < 0 or not mapping:
            continue

        rows = list(ws.iter_rows(values_only=True))
        for row in rows[data_start:]:
            item = _parse_row_to_item(list(row), mapping, meta)
            if item:
                all_items.append(item)

    # Строим текст для AI fallback
    text_lines = []
    if meta_global.get("contractor"):
        text_lines.append(f"Подрядчик: {meta_global['contractor']}")
    if meta_global.get("project"):
        text_lines.append(f"Объект: {meta_global['project']}")
    for item in all_items:
        parts = [item["name"]]
        if item.get("quantity"):   parts.append(f"кол-во: {item['quantity']}")
        if item.get("unit"):       parts.append(item["unit"])
        if item.get("unit_price"): parts.append(f"цена: {item['unit_price']}")
        if item.get("total_price"):parts.append(f"сумма: {item['total_price']}")
        text_lines.append(" | ".join(parts))

    return {
        "success":    True,
        "extension":  ".xlsx",
        "file":       file_path.name,
        "text":       "\n".join(text_lines),
        "items":      all_items,
        "metadata":   meta_global,
        "characters": len("\n".join(text_lines)),
    }


# ─── PDF парсер ───────────────────────────────────────────────────

def _parse_pdf_tables(file_path: Path) -> tuple[list[dict], str]:
    """Извлекает таблицы и текст из PDF через pdfplumber."""
    items    = []
    texts    = []
    metadata = {"contractor": "", "date": ""}

    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            if page_text:
                texts.append(f"[Стр.{page_num+1}]\n{page_text}")

            # Пробуем таблицы
            tables = page.extract_tables()
            for table in tables:
                if not table or len(table) < 2:
                    continue

                # Ищем заголовок таблицы
                header_rows = []
                data_start  = 0
                for ri, row in enumerate(table[:8]):
                    texts_in_row = [_cell_text(c) for c in (row or []) if c]
                    has_num_marker = any(
                        t.strip() in ("1", "2", "3") for t in texts_in_row
                    )
                    if has_num_marker:
                        data_start = ri + 1
                        break
                    has_header_kw = any(
                        any(kw in t for kw in ["наименование", "количество", "стоимость", "цена"])
                        for t in texts_in_row
                    )
                    if has_header_kw:
                        header_rows.append(row)
                        data_start = ri + 1

                if not header_rows:
                    header_rows = table[:1]
                    data_start  = 1

                mapping = _detect_column_mapping(header_rows)
                if not mapping:
                    # Числовая инференция
                    mapping = _infer_mapping_by_numbers(table[data_start:])

                for row in table[data_start:]:
                    if not row:
                        continue
                    item = _parse_row_to_item(list(row), mapping, metadata)
                    if item:
                        items.append(item)

    return items, "\n\n".join(texts)


def _infer_mapping_by_numbers(rows: list[list]) -> dict:
    """Числовая инференция маппинга: quantity × unit_price ≈ total_price."""
    if not rows:
        return {}

    n_cols = max((len(r) for r in rows if r), default=0)
    if n_cols < 2:
        return {}

    col_vals = [[] for _ in range(n_cols)]
    for row in rows:
        for ci, cell in enumerate(row or []):
            v = _to_float(cell)
            if v is not None and v > 0:
                col_vals[ci].append(v)

    numeric_cols = [i for i in range(n_cols) if len(col_vals[i]) >= max(1, len(rows) // 3)]
    if len(numeric_cols) < 2:
        return {}

    best_mapping = {}
    best_score   = 0

    for qi in numeric_cols:
        for ui in numeric_cols:
            if ui == qi:
                continue
            for ti in numeric_cols:
                if ti in (qi, ui):
                    continue
                matches = 0
                for row in rows:
                    if not row:
                        continue
                    q = _to_float(row[qi] if qi < len(row) else None)
                    u = _to_float(row[ui] if ui < len(row) else None)
                    t = _to_float(row[ti] if ti < len(row) else None)
                    if q and u and t and t > 0:
                        if abs(q * u - t) / t < 0.03:
                            matches += 1
                if matches > best_score:
                    best_score   = matches
                    best_mapping = {"quantity": qi, "unit_price": ui, "total_price": ti}

    if not best_mapping and len(numeric_cols) >= 2:
        sums = sorted([(i, sum(col_vals[i]) / len(col_vals[i])) for i in numeric_cols], key=lambda x: x[1])
        best_mapping = {"unit_price": sums[0][0], "total_price": sums[-1][0]}

    # Первая нечисловая колонка → name
    for i in range(n_cols):
        if i not in best_mapping.values():
            sample = [str(r[i]) for r in rows[:5] if r and i < len(r) and r[i]]
            if any(len(s) > 3 and not _to_float(s) for s in sample):
                best_mapping["name"] = i
                break

    return best_mapping


def parse_pdf(file_path: Path) -> dict:
    items, text = _parse_pdf_tables(file_path)
    return {
        "success":    True,
        "extension":  ".pdf",
        "file":       file_path.name,
        "text":       text,
        "items":      items,
        "metadata":   {},
        "characters": len(text),
    }


# ─── DOCX парсер ──────────────────────────────────────────────────

def parse_docx(file_path: Path) -> dict:
    doc   = docx.Document(file_path)
    items = []
    texts = []
    metadata = {"contractor": "", "date": ""}

    # Параграфы — метаданные
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        texts.append(t)
        tl = t.lower()
        for meta_key, keywords in _META_KEYWORDS.items():
            if any(kw in tl for kw in keywords) and ":" in t:
                val = t.split(":", 1)[-1].strip()
                if val and len(val) > 3:
                    if meta_key == "contractor" and not metadata["contractor"]:
                        metadata["contractor"] = val

    # Таблицы
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])

        if len(rows) < 2:
            continue

        # Заголовок
        header_rows = []
        data_start  = 0
        for ri, row in enumerate(rows[:8]):
            has_num = any(c.strip() in ("1", "2", "3") for c in row)
            if has_num:
                data_start = ri + 1
                break
            has_kw = any(
                any(kw in c.lower() for kw in ["наименование", "количество", "стоимость", "цена"])
                for c in row
            )
            if has_kw:
                header_rows.append(row)
                data_start = ri + 1

        if not header_rows:
            header_rows = rows[:1]
            data_start  = 1

        mapping = _detect_column_mapping(header_rows)
        if not mapping:
            mapping = _infer_mapping_by_numbers(rows[data_start:])

        for row in rows[data_start:]:
            item = _parse_row_to_item(row, mapping, metadata)
            if item:
                items.append(item)
                texts.append(" | ".join(str(v) for v in row if v))

    return {
        "success":    True,
        "extension":  ".docx",
        "file":       file_path.name,
        "text":       "\n".join(texts),
        "items":      items,
        "metadata":   metadata,
        "characters": len("\n".join(texts)),
    }


# ─── Публичный API ────────────────────────────────────────────────

def parse_file(file_path: str | Path) -> dict:
    file_path = Path(file_path)
    if not file_path.exists():
        return {"success": False, "file": file_path.name, "error": "Файл не найден", "text": "", "items": []}

    parsers = {".pdf": parse_pdf, ".docx": parse_docx, ".xlsx": parse_xlsx}
    ext = file_path.suffix.lower()

    if ext not in parsers:
        return {"success": False, "file": file_path.name, "error": f"Неподдерживаемый формат: {ext}", "text": "", "items": []}

    try:
        return parsers[ext](file_path)
    except Exception as e:
        return {"success": False, "file": file_path.name, "error": str(e), "text": "", "items": []}